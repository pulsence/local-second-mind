from __future__ import annotations

from pathlib import Path

from docx import Document

from lsm.utils.file_graph import get_file_graph


def _find_node(graph, node_type: str, name: str):
    for node in graph.nodes:
        if node.node_type == node_type and node.name == name:
            return node
    raise AssertionError(f"Missing node {node_type}:{name}")


def _find_node_by_line(graph, node_type: str, start_line: int):
    for node in graph.nodes:
        if node.node_type == node_type and node.start_line == start_line:
            return node
    raise AssertionError(f"Missing node {node_type} at line {start_line}")


def test_text_graph_markdown_headings_and_lists(tmp_path: Path) -> None:
    text = (
        "# Title\n\n"
        "Intro paragraph.\n\n"
        "## Section A\n\n"
        "Paragraph A.\n\n"
        "- Item 1\n"
        "- Item 2\n\n"
        "## Section B\n\n"
        "Paragraph B.\n"
    )
    path = tmp_path / "sample.md"
    path.write_text(text, encoding="utf-8")

    graph = get_file_graph(path)

    title = _find_node(graph, "heading", "Title")
    section_a = _find_node(graph, "heading", "Section A")
    section_b = _find_node(graph, "heading", "Section B")

    intro = _find_node_by_line(graph, "paragraph", 3)
    para_a = _find_node_by_line(graph, "paragraph", 7)
    list_block = _find_node_by_line(graph, "list", 9)
    para_b = _find_node_by_line(graph, "paragraph", 14)

    assert title.start_line == 1
    assert title.end_line == 14
    assert section_a.start_line == 5
    assert section_b.start_line == 12
    assert section_a.end_line == 10
    assert section_b.end_line == 14

    assert intro.parent_id == title.id
    assert section_a.parent_id == title.id
    assert para_a.parent_id == section_a.id
    assert list_block.parent_id == section_a.id
    assert section_b.parent_id == title.id
    assert para_b.parent_id == section_b.id

    assert list_block.end_line == 10


def test_text_graph_plain_text_headings(tmp_path: Path) -> None:
    text = (
        "Section 1: Overview\n\n"
        "Overview paragraph.\n\n"
        "Section 1.1: Details\n\n"
        "Detail paragraph.\n"
    )
    path = tmp_path / "plain.txt"
    path.write_text(text, encoding="utf-8")

    graph = get_file_graph(path)

    overview = _find_node(graph, "heading", "Section 1: Overview")
    details = _find_node(graph, "heading", "Section 1.1: Details")

    assert overview.start_line == 1
    assert details.start_line == 5
    assert details.parent_id == overview.id
    assert overview.end_line == 7
    assert details.end_line == 7


def test_text_graph_docx_headings(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("Doc Title", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_heading("Doc Section", level=2)
    doc.add_paragraph("Body paragraph.")

    path = tmp_path / "sample.docx"
    doc.save(path)

    graph = get_file_graph(path)

    title = _find_node(graph, "heading", "Doc Title")
    section = _find_node(graph, "heading", "Doc Section")

    assert title.metadata.get("level") == 1
    assert section.metadata.get("level") == 2
    assert section.parent_id == title.id
    assert title.end_line == 7
    assert section.end_line == 7
